"""
01 Document parser (module 02).

PDF / DOCX / TXT extraction, cleaning, optional traditional-to-simplified Chinese, and quality gating for chunks.
"""

import re
import hashlib
from pathlib import Path
from typing import Tuple, List, Set, Dict, Optional
import logging

try:
    from simhash import Simhash
    SIMHASH_AVAILABLE = True
except ImportError:
    SIMHASH_AVAILABLE = False
    logging.warning("simhash not available, paragraph-level deduplication will be disabled")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not available, PDF parsing will be limited")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available, DOCX parsing will be limited")

try:
    import opencc
    OPENC_AVAILABLE = True
except ImportError:
    OPENC_AVAILABLE = False
    logging.warning("opencc not available, traditional Chinese conversion will be skipped")

logger = logging.getLogger(__name__)


class DocumentParser:
    """Document parser for PDF, DOCX, TXT."""

    def __init__(self, convert_traditional: bool = True, chinese_only: bool = True, min_chinese_ratio: float = 0.60):
        """
        Args:
            convert_traditional: Traditional to simplified Chinese
            chinese_only: Keep only Chinese-dominant content
            min_chinese_ratio: Minimum Chinese character ratio for language detection
        """
        self.convert_traditional = convert_traditional
        self.chinese_only = chinese_only
        self.min_chinese_ratio = min_chinese_ratio
        if convert_traditional and OPENC_AVAILABLE:
            try:
                self.converter = opencc.OpenCC("t2s")
            except Exception as e:
                logger.warning(
                    f"Failed to initialize opencc converter: {e}. Traditional Chinese conversion will be disabled."
                )
                self.converter = None
        else:
            self.converter = None

    @staticmethod
    def _count_replacement(text: str) -> int:
        return text.count("\uFFFD")

    @staticmethod
    def _count_extb(text: str) -> int:
        return len(re.findall(r"[\U00020000-\U0002A6DF]", text))

    @staticmethod
    def fix_encoding_issues(text: str) -> str:
        """Fix encoding issues and garbled characters. Detect/remove garbled sequences, try to fix encoding errors."""
        if not text:
            return ""
        lines = text.split('\n')
        fixed_lines = []
        
        for line in lines:
            # Compute ratio of valid characters in the line
            total_chars = len(line)
            if total_chars == 0:
                fixed_lines.append(line)
                continue
            
            # Count valid characters
            chinese_chars = len(re.findall(r'[\u4e00-\u9fa5]', line))
            english_chars = len(re.findall(r'[a-zA-Z]', line))
            digits = len(re.findall(r'\d', line))
            common_punct = len(re.findall(r'[，。！？、；：""''（）【】《》·—\s,.!?;:()\[\]<>/\\\-]', line))
            valid_chars = chinese_chars + english_chars + digits + common_punct
            
            # If valid character ratio is too low (<30%), likely garbled; try to fix or skip
            if total_chars > 10 and valid_chars / total_chars < 0.30:
                # Try to extract plausible Chinese/English/digit fragments
                chinese_pattern = re.findall(r'[\u4e00-\u9fa5]+', line)
                english_pattern = re.findall(r'[a-zA-Z]+', line)
                digit_pattern = re.findall(r'\d+', line)
                if chinese_pattern or (english_pattern and len(' '.join(english_pattern)) > 20):
                    fixed_line = ' '.join(chinese_pattern + english_pattern + digit_pattern)
                    if len(fixed_line.strip()) > 5:
                        fixed_lines.append(fixed_line)
            else:
                # Normal line; clean encoding-error characters (PUA, replacement, control chars)
                line = re.sub(r'[\uE000-\uF8FF]', '', line)
                line = line.replace('\uFFFD', '')
                line = re.sub(r'[\x00-\x08\x0B-\x1F\x7F]', '', line)
                fixed_lines.append(line)
        
        fixed_text = '\n'.join(fixed_lines)
        
        # Remove obvious garbled paragraphs (high ratio of non-standard characters)
        paragraphs = fixed_text.split('\n\n')
        fixed_paragraphs = []
        
        for para in paragraphs:
            if not para.strip():
                fixed_paragraphs.append(para)
                continue
            
            total = len(para)
            chinese = len(re.findall(r'[\u4e00-\u9fa5]', para))
            english = len(re.findall(r'[a-zA-Z]', para))
            digits = len(re.findall(r'\d', para))
            common_punct = len(re.findall(r'[，。！？、；：""''（）【】《》·—\s,.!?;:()\[\]<>/\\\-]', para))
            valid = chinese + english + digits + common_punct
            
            if total > 0 and valid / total >= 0.50:
                fixed_paragraphs.append(para)
            elif total < 20 and (chinese > 0 or english > 5):
                fixed_paragraphs.append(para)
        
        return '\n\n'.join(fixed_paragraphs)

    @staticmethod
    def sanitize_text(text: str) -> str:
        if not text:
            return ""
        # Private Use Area (font mapping residue) and replacement character
        text = re.sub(r"[\uE000-\uF8FF]", "", text)
        text = text.replace("\uFFFD", "")
        # Control characters (keep \n \t)
        text = re.sub(r"[\x00-\x08\x0B-\x1F\x7F]", "", text)
        # Normalize whitespace
        text = re.sub(r"[ \u3000]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def is_valid_text(
        self,
        text: str,
        min_len: int = 80,
        min_valid_ratio: float = 0.70,
        max_uncommon_ratio: float = 0.20,
        max_extb_ratio: float = 0.01,
        max_extb_abs: int = 10,
    ) -> Tuple[bool, str]:
        if not text or len(text.strip()) < min_len:
            return False, "Text too short or empty"

        total = len(text)
        if total == 0:
            return False, "Text is empty"

        # Replacement character is almost always a bad signal
        if self._count_replacement(text) > 0:
            return False, "Contains replacement character"

        extb = self._count_extb(text)
        if extb >= max_extb_abs or (extb / total) > max_extb_ratio:
            return False, "Too many extended-B characters"

        common_chinese = len(re.findall(r"[\u4e00-\u9fa5]", text))
        english = len(re.findall(r"[a-zA-Z]", text))
        digits = len(re.findall(r"\d", text))
        punctuation = len(re.findall(r"[，。！？、；：\"\"''（）【】《》·—\s,.!?;:()\[\]<>/\\\-]", text))

        valid = common_chinese + english + digits + punctuation
        valid_ratio = valid / total
        if valid_ratio < min_valid_ratio:
            return False, f"Valid character ratio too low ({valid_ratio:.2%})"

        uncommon = len(
            re.findall(r"[^\u0000-\u007F\u4e00-\u9fa5\s，。！？、；：\"\"''（）【】《》·—,.!?;:()\[\]<>/\\\-]", text)
        )
        uncommon_ratio = uncommon / total
        if uncommon_ratio > max_uncommon_ratio:
            return False, f"Uncommon character ratio too high ({uncommon_ratio:.2%})"

        chinese_ratio = common_chinese / total if total > 0 else 0
        english_ratio = english / total if total > 0 else 0
        
        if self.chinese_only:
            if chinese_ratio < self.min_chinese_ratio:
                return False, f"Chinese ratio too low ({chinese_ratio:.2%}), required ≥{self.min_chinese_ratio:.2%}"
        else:
            if chinese_ratio < 0.10 and english_ratio < 0.10:
                return False, "Neither predominantly Chinese nor predominantly English"

        return True, "OK"

    def parse_pdf(self, file_path: str) -> str:
        if not PDFPLUMBER_AVAILABLE:
            raise ImportError("pdfplumber is required for PDF parsing. Install it with: pip install pdfplumber")

        text_parts = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text() or ""
                        # Fix encoding issues first (handle possible garbled text)
                        page_text = self.fix_encoding_issues(page_text)
                        page_text = self.clean_text(page_text)
                        if self.convert_traditional:
                            page_text = self.convert_to_simplified(page_text)
                        page_text = self.sanitize_text(page_text)
                        ok, _ = self.is_valid_text(page_text, min_len=80)
                        if ok:
                            text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num} of {file_path}: {e}")
                        continue
        except Exception as e:
            logger.error(f"Error opening PDF file {file_path}: {e}")
            raise

        return "\n".join(text_parts)

    def parse_docx(self, file_path: str) -> str:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing. Install it with: pip install python-docx")

        try:
            doc = Document(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)

            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {e}")
            raise

    def parse_txt(self, file_path: str, encoding: str = "utf-8") -> str:
        """
        Parse TXT file with auto-detection of encoding and text quality validation.

        Key issue: reading a GBK file as UTF-8 does not raise UnicodeDecodeError but yields garbled text.
        Approach: after reading, validate text quality; if garbled, try other encodings.
        """
        encodings = [encoding, "utf-8", "gbk", "gb2312", "gb18030", "big5", "latin1"]
        best_text = None
        best_score = -1
        best_encoding = None
        
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc, errors='replace') as f:
                    text = f.read()
                    
                    if not text:
                        continue
                    
                    total_chars = len(text)
                    if total_chars == 0:
                        continue
                    
                    chinese_chars = len(re.findall(r'[\u4e00-\u9fa5]', text))
                    english_chars = len(re.findall(r'[a-zA-Z]', text))
                    digits = len(re.findall(r'\d', text))
                    common_punct = len(re.findall(r'[，。！？、；：""''（）【】《》·—\s,.!?;:()\[\]<>/\\\-]', text))
                    valid_chars = chinese_chars + english_chars + digits + common_punct
                    
                    valid_ratio = valid_chars / total_chars if total_chars > 0 else 0
                    chinese_ratio = chinese_chars / total_chars if total_chars > 0 else 0
                    
                    if self.chinese_only:
                        score = valid_ratio * 0.6 + chinese_ratio * 0.4
                    else:
                        score = valid_ratio
                    
                    replacement_count = text.count('\uFFFD')
                    if replacement_count > total_chars * 0.1:
                        score *= 0.3
                    
                    private_chars = len(re.findall(r'[\uE000-\uF8FF]', text))
                    if private_chars > total_chars * 0.05:
                        score *= 0.5
                    
                    if valid_ratio < 0.3:
                        score *= 0.2
                    
                    if score > best_score:
                        best_score = score
                        best_text = text
                        best_encoding = enc
                    
                    if score > 0.8:
                        logger.debug(f"TXT file {file_path} using encoding {enc}, quality score: {score:.2f}")
                        return text
                        
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(f"Error reading TXT file {file_path} with encoding {enc}: {e}")
                continue
        
        if best_text is not None and best_score > 0.1:
            logger.info(f"TXT file {file_path} using encoding {best_encoding}, quality score: {best_score:.2f}")
            return best_text
        
        try:
            import chardet
            with open(file_path, "rb") as f:
                raw_data = f.read()
                detected = chardet.detect(raw_data)
                if detected and detected['encoding']:
                    detected_enc = detected['encoding']
                    confidence = detected.get('confidence', 0)
                    if confidence > 0.7:
                        logger.info(f"chardet detected encoding: {detected_enc} (confidence: {confidence:.2f})")
                        with open(file_path, "r", encoding=detected_enc, errors='replace') as f:
                            return f.read()
        except ImportError:
            logger.debug("chardet not installed, skipping auto encoding detection")
        except Exception as e:
            logger.warning(f"chardet encoding detection failed: {e}")
        
        if best_text is None:
            logger.warning(f"Cannot determine correct encoding for TXT file {file_path}, using UTF-8 with errors='replace'")
            with open(file_path, "r", encoding="utf-8", errors='replace') as f:
                return f.read()
        
        return best_text

    def parse_file(self, file_path: str) -> str:
        file_path_obj = Path(file_path)
        suffix = file_path_obj.suffix.lower()
        if suffix == ".pdf":
            return self.parse_pdf(str(file_path_obj))
        if suffix == ".docx":
            return self.parse_docx(str(file_path_obj))
        if suffix == ".txt":
            return self.parse_txt(str(file_path_obj))
        raise ValueError(f"Unsupported file format: {suffix}")

    @staticmethod
    def _is_structural_noise(text: str) -> bool:
        """Whether the text is structural noise (TOC, disclaimer, header/footer templates, etc.)."""
        text_lower = text.lower()
        # TOC: many short lines + numeric/alphabetic numbering
        if len(text.split('\n')) > 20 and re.search(r'^\s*[0-9a-z\.\-]+\s+', text, re.MULTILINE):
            if len([l for l in text.split('\n') if len(l.strip()) < 30]) / max(len(text.split('\n')), 1) > 0.7:
                return True
        # Disclaimer/copyright page (keywords include Chinese and English)
        noise_keywords = ['免责声明', '版权', 'copyright', 'disclaimer', '版权所有', '本报告', '本文件仅供']
        if any(kw in text_lower for kw in noise_keywords) and len(text) < 500:
            return True
        # Header/footer template (repeated)
        lines = text.split('\n')
        if len(lines) > 5:
            first_line = lines[0].strip()
            if first_line and lines.count(first_line) >= len(lines) * 0.3:
                return True
        return False

    def clean_text(self, text: str) -> str:
        """Clean text: remove headers, footers, page numbers, and structural noise."""
        text = re.sub(r"^\d+$", "", text, flags=re.MULTILINE)
        text = re.sub(r"第\s*\d+\s*页\s*[/\\]\s*共\s*\d+\s*页", "", text, flags=re.IGNORECASE)
        text = re.sub(r"Page\s+\d+\s+of\s+\d+", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\[图片\]|\[图\d+\]|\[Image\]|\[Figure\s+\d+\]", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s{3,}", " ", text)
        text = "\n".join(line.strip() for line in text.split("\n"))
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def deduplicate_lines(text: str) -> Tuple[str, int]:
        """Line-level deduplication (by hash); returns deduplicated text and number of lines removed."""
        lines = text.split('\n')
        seen_hashes: Set[str] = set()
        unique_lines: List[str] = []
        removed_count = 0
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                unique_lines.append(line)
                continue
            line_hash = hashlib.md5(line_stripped.encode('utf-8')).hexdigest()
            if line_hash not in seen_hashes:
                seen_hashes.add(line_hash)
                unique_lines.append(line)
            else:
                removed_count += 1
        
        return '\n'.join(unique_lines), removed_count

    @staticmethod
    def deduplicate_paragraphs(text: str, simhash_threshold: int = 3) -> Tuple[str, int]:
        """Paragraph-level deduplication (simhash); returns deduplicated text and number of paragraphs removed."""
        if not SIMHASH_AVAILABLE:
            return text, 0
        
        paragraphs = re.split(r'\n\s*\n', text)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]
        if len(paragraphs) <= 1:
            return text, 0
        
        seen_hashes: List[int] = []
        unique_paragraphs: List[str] = []
        removed_count = 0
        
        for para in paragraphs:
            if len(para) < 20:  # Skip simhash for very short paragraphs
                unique_paragraphs.append(para)
                continue
            
            try:
                para_hash = Simhash(para).value
                is_duplicate = False
                for seen_hash in seen_hashes:
                    hamming = bin(para_hash ^ seen_hash).count('1')
                    if hamming <= simhash_threshold:
                        is_duplicate = True
                        removed_count += 1
                        break
                if not is_duplicate:
                    seen_hashes.append(para_hash)
                    unique_paragraphs.append(para)
            except Exception as e:
                logger.warning(f"SimHash calculation failed for paragraph: {e}")
                unique_paragraphs.append(para)
        
        return '\n\n'.join(unique_paragraphs), removed_count

    def convert_to_simplified(self, text: str) -> str:
        if self.converter:
            try:
                return self.converter.convert(text)
            except Exception as e:
                logger.warning(f"Error converting traditional to simplified: {e}")
                return text
        return text

    def parse_and_clean(
        self, 
        file_path: str, 
        enable_line_dedup: bool = True,
        enable_para_dedup: bool = True,
        remove_structural_noise: bool = True
    ) -> Tuple[str, Dict[str, int]]:
        """
        Parse and clean file; return cleaned text and stats.

        Returns:
            (cleaned_text, stats_dict)
            stats_dict keys: line_dedup_removed, para_dedup_removed, structural_noise_removed
        """
        text = self.parse_file(file_path)
        stats: Dict[str, int] = {"line_dedup_removed": 0, "para_dedup_removed": 0, "structural_noise_removed": 0}
        
        text = self.clean_text(text)
        
        if remove_structural_noise:
            paragraphs = re.split(r'\n\s*\n', text)
            filtered_paragraphs = []
            for para in paragraphs:
                if not self._is_structural_noise(para.strip()):
                    filtered_paragraphs.append(para)
                else:
                    stats["structural_noise_removed"] += 1
            text = '\n\n'.join(filtered_paragraphs)
        
        if self.convert_traditional:
            text = self.convert_to_simplified(text)
        
        if enable_line_dedup:
            text, removed = self.deduplicate_lines(text)
            stats["line_dedup_removed"] = removed
        
        if enable_para_dedup:
            text, removed = self.deduplicate_paragraphs(text)
            stats["para_dedup_removed"] = removed
        
        text = self.sanitize_text(text)
        
        return text, stats


